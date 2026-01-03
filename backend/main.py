from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import io
import json

app = FastAPI(title="Docx Editor Backend")

# --------------------------------------------------
# CORS
# --------------------------------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080",
                 "https://word-doc-editor-p.vercel.app",
                 "https://word-doc-editor-psi.vercel.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --------------------------------------------------
# SAFE paragraph replacement (NO duplication)
# --------------------------------------------------
def replace_paragraph_text(paragraph, replacements: dict):
    original_text = paragraph.text.strip()

    if original_text in replacements:
        new_text = replacements[original_text]

        # Preserve paragraph formatting, replace content once
        paragraph.clear()
        paragraph.add_run(new_text)

# --------------------------------------------------
# EDIT DOCX
# --------------------------------------------------
@app.post("/edit-docx")
async def edit_docx(
    file: UploadFile = File(...),
    text_map: str = Form(...)
):
    contents = await file.read()
    doc = Document(io.BytesIO(contents))

    replacements = json.loads(text_map)

    # Normal paragraphs
    for paragraph in doc.paragraphs:
        replace_paragraph_text(paragraph, replacements)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, replacements)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=edited_{file.filename}"
        },
    )

# --------------------------------------------------
# EXTRACT TEXT (DISPLAY ONLY)
# --------------------------------------------------
@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    contents = await file.read()
    doc = Document(io.BytesIO(contents))

    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())

    return JSONResponse({
        "text": "\n\n".join(paragraphs)
    })
